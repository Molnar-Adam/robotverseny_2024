#!/usr/bin/env python3
"""Generate a Word document about ROS1 vs ROS2 test differences.

The script scans the current repository test layout and creates a .docx
document that summarizes:
- current test inventory
- key ROS1 vs ROS2 test framework differences
- concrete migration impact for this codebase
- suggested migration checklist
"""

from __future__ import annotations

import argparse
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List, Sequence


try:
    from docx import Document
except ImportError as exc:  # pragma: no cover - import guard
    raise SystemExit(
        "Missing dependency: python-docx. Install with: pip install python-docx"
    ) from exc


@dataclass
class TestInventory:
    python_unit_tests: List[Path]
    cpp_unit_tests: List[Path]
    integration_py_tests: List[Path]
    integration_launch_tests: List[Path]

    @property
    def total(self) -> int:
        return (
            len(self.python_unit_tests)
            + len(self.cpp_unit_tests)
            + len(self.integration_py_tests)
            + len(self.integration_launch_tests)
        )


def _relative_paths(paths: Sequence[Path], root: Path) -> List[str]:
    return sorted(str(path.relative_to(root)).replace("\\", "/") for path in paths)


def collect_test_inventory(repo_root: Path) -> TestInventory:
    tests_root = repo_root / "tests"
    integration_root = tests_root / "integration"

    python_unit = [
        p
        for p in tests_root.glob("test_*.py")
        if p.is_file() and p.parent == tests_root
    ]
    cpp_unit = [p for p in tests_root.glob("*.cpp") if p.is_file()]
    integration_py = [p for p in integration_root.glob("test_*.py") if p.is_file()]
    integration_launch = [p for p in integration_root.glob("*.test") if p.is_file()]

    return TestInventory(
        python_unit_tests=sorted(python_unit),
        cpp_unit_tests=sorted(cpp_unit),
        integration_py_tests=sorted(integration_py),
        integration_launch_tests=sorted(integration_launch),
    )


def _add_bullet_list(document: Document, title: str, items: Sequence[str]) -> None:
    document.add_paragraph(title)
    if not items:
        document.add_paragraph("- (nincs ilyen teszt)")
        return
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def build_document(repo_root: Path, inventory: TestInventory) -> Document:
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    document = Document()

    document.add_heading("ROS1 vs ROS2 tesztkulonbsegek", level=0)
    document.add_paragraph(
        "Automatikusan generalt dokumentacio a jelenlegi tesztkeszlet alapjan."
    )
    document.add_paragraph(f"Generalas ideje: {now}")
    document.add_paragraph(f"Projekt gyokermappa: {repo_root}")

    document.add_heading("1. Jelenlegi tesztallapot", level=1)
    document.add_paragraph(f"Osszes talalt tesztfajl: {inventory.total}")
    document.add_paragraph(
        "Megjegyzes: az osszesites fajlszamu nezet, nem kulonall tesztesetszamu nezet."
    )

    _add_bullet_list(
        document,
        f"Python unit tesztfajlok ({len(inventory.python_unit_tests)}):",
        _relative_paths(inventory.python_unit_tests, repo_root),
    )
    _add_bullet_list(
        document,
        f"C++ unit tesztfajlok ({len(inventory.cpp_unit_tests)}):",
        _relative_paths(inventory.cpp_unit_tests, repo_root),
    )
    _add_bullet_list(
        document,
        f"Integracios Python tesztfajlok ({len(inventory.integration_py_tests)}):",
        _relative_paths(inventory.integration_py_tests, repo_root),
    )
    _add_bullet_list(
        document,
        f"Integracios launch/rostest fajlok ({len(inventory.integration_launch_tests)}):",
        _relative_paths(inventory.integration_launch_tests, repo_root),
    )

    document.add_heading("2. Kulonbsegek ROS1 es ROS2 kozott", level=1)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Terulet"
    header_cells[1].text = "Jelenlegi ROS1 allapot"
    header_cells[2].text = "ROS2 valtozat"
    header_cells[3].text = "Varhato projekt-hatas"

    differences = [
        (
            "Integracios teszt launcher",
            "rostest + .test fajlok",
            "pytest + launch_testing (Python launch API)",
            "A .test fajlokat launch.py + pytest alapu tesztekre kell cserelni.",
        ),
        (
            "Build/test parancsok",
            "catkin_make run_tests, catkin_test_results",
            "colcon test, colcon test-result --verbose",
            "CI es lokalis parancslista atalakitasa szukseges.",
        ),
        (
            "Python API",
            "rospy",
            "rclpy",
            "Node letrehozas, spineles, parameterek es ido API atiras kell.",
        ),
        (
            "C++ API",
            "roscpp",
            "rclcpp",
            "Publisher/subscriber es QoS kezeles atalakul.",
        ),
        (
            "Uzenet atvitel",
            "ROS1 topic szemantika",
            "DDS + QoS profilok",
            "Tesztekben explicit QoS beallitas kell a stabil eredmenyhez.",
        ),
        (
            "Parameterezes",
            "rosparam + launch param",
            "declare_parameter/get_parameter + launch substitutions",
            "Node startolaskor explicit parameter deklaracio szukseges.",
        ),
        (
            "Launch rendszer",
            "XML launch",
            "Python launch (launch, launch_ros)",
            "Integracios tesztek launch resze Pythonra koltozik.",
        ),
        (
            "Csomag metadata",
            "catkin package.xml + CMakeLists.txt",
            "ament_cmake vagy ament_python",
            "Teszt targeteket es fuggosegeket ujra kell definialni.",
        ),
    ]

    for area, ros1_now, ros2_way, impact in differences:
        row = table.add_row().cells
        row[0].text = area
        row[1].text = ros1_now
        row[2].text = ros2_way
        row[3].text = impact

    document.add_heading("3. Konret valtozasok a jelenlegi tesztekre", level=1)
    concrete_changes = [
        "tests/integration/*.test fajlok -> ROS2 launch_testing alapra atirva.",
        "Python node tesztekben rospy importok -> rclpy importokra cserelve.",
        "Topic varakozasnal QoS profilok explicit beallitasa (pl. SensorDataQoS).",
        "catkin_make run_tests helyett colcon test futtatasi lepes.",
        "C++ unit teszt targetek catkin_add_gtest helyett ament_add_gtest alapon.",
        "Szukseg esetben ros1_bridge/kompatibilitasi tesztek a fokozatos atallasra.",
    ]
    for item in concrete_changes:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("4. Javasolt migracios terv", level=1)
    migration_steps = [
        "Fuggosegterkep keszitese: minden ROS1 csomag es API hivatkozas listazasa.",
        "Build rendszer dontes: ament_cmake es/vagy ament_python.",
        "Node-ok portolasa (rospy/roscpp -> rclpy/rclcpp).",
        "Integracios tesztek launch_testing alapra koltoztetese.",
        "QoS profilok finomhangolasa a stabil tesztfutasokhoz.",
        "CI pipeline frissitese colcon test parancsokra.",
        "Vegso regresszios kor ROS2 alatt (unit + integracios + rendszer teszt).",
    ]
    for index, step in enumerate(migration_steps, start=1):
        document.add_paragraph(f"{index}. {step}")

    document.add_heading("5. Mintaparancsok ROS2-ben", level=1)
    cmd_samples = [
        "colcon build --packages-select megoldas_gyor24",
        "colcon test --packages-select megoldas_gyor24",
        "colcon test-result --verbose",
        "pytest -q",
    ]
    for cmd in cmd_samples:
        document.add_paragraph(cmd, style="List Bullet")

    return document


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Word dokumentum generalasa a jelenlegi tesztek es egy ROS2-es valtozat "
            "kulonbsegeirol."
        )
    )
    parser.add_argument(
        "--repo-root",
        type=Path,
        default=Path(__file__).resolve().parents[1],
        help="Projekt gyokermappa (default: script alapjan automatikus).",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("docs/ros1_vs_ros2_tesztkulonbsegek.docx"),
        help="Kimeneti .docx fajl utvonala.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    repo_root = args.repo_root.resolve()
    output_path = args.output
    if not output_path.is_absolute():
        output_path = repo_root / output_path

    inventory = collect_test_inventory(repo_root)
    document = build_document(repo_root, inventory)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)

    print(f"Dokumentum elkeszult: {output_path}")
    print(f"Talalt tesztfajlok osszesen: {inventory.total}")


if __name__ == "__main__":
    main()
